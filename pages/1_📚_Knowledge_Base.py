import uuid

import streamlit as st
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

from src.retrieval.vector_store import get_embeddings

st.set_page_config(page_title="User Document Upload", layout="wide")

# File uploader
uploaded_files = st.file_uploader(
    "Upload text, PDF, or DOCX files", type=["txt", "pdf", "docx"], accept_multiple_files=True
)

if uploaded_files:
    session_id = str(uuid.uuid4())
    documents = []

    for file in uploaded_files:
        if file.name.endswith(".txt"):
            content = file.read().decode("utf-8")
        elif file.name.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(file)
            content = "\n".join([page.extract_text() for page in reader.pages])
        elif file.name.endswith(".docx"):
            from docx import Document as DocxDocument

            doc = DocxDocument(file)
            content = "\n".join([para.text for para in doc.paragraphs])
        else:
            continue

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(content)
        docs = [
            Document(page_content=chunk, metadata={"source": "user_upload"})
            for chunk in chunks
        ]
        documents.extend(docs)

    # Store session_id in st.session_state so other pages can access it
    st.session_state.upload_session_id = session_id

    # Add to session-scoped collection (persistent to disk, not in-memory)
    embedding = get_embeddings()
    vs = Chroma(
        collection_name=f"session_{session_id}",
        embedding_function=embedding,
        persist_directory=f"data/index/session_uploads/{session_id}",
    )
    vs.add_documents(documents)
    st.success("Documents uploaded and indexed for this session!")